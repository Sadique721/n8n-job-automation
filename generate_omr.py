import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """
    Set cell padding in dxa (1 pt = 20 dxa)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_omr_sheet():
    doc = Document()
    
    # Set page size to A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 width
    section.page_height = Inches(11.69) # A4 height
    
    # Narrow margins to fit everything on a single page
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    
    # Theme color: Professional Crimson Red (Hex: A6192E)
    theme_color = RGBColor(166, 25, 46)
    theme_hex = "A6192E"
    
    # Use the first paragraph for the Title (avoiding default empty paragraph bug)
    if len(doc.paragraphs) > 0:
        title_p = doc.paragraphs[0]
    else:
        title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("OMR ANSWER SHEET")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = theme_color
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    
    # Header Info Table (Name, Date, TA, Right, Wrong, Total Marks)
    # 3 rows, 4 columns: Label1, Value1, Label2, Value2
    info_table = doc.add_table(rows=3, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    
    # Set column widths for info table
    # Total width = 7.47 inches (8.27 - 0.8)
    widths = [Inches(1.2), Inches(2.5), Inches(1.2), Inches(2.57)]
    for row in info_table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            set_cell_margins(row.cells[idx], top=40, bottom=40, left=40, right=40)
            
    # Labels and fields structure matching the user's image
    headers = [
        [("Name:", True), ("___________________________", False), ("Date:", True), ("_________________", False)],
        [("TA:", True), ("_________________", False), ("Right:", True), ("_________________", False)],
        [("Wrong:", True), ("_________________", False), ("Total Marks:", True), ("_________________", False)]
    ]
    
    # Populate Header Info
    for r_idx, row_data in enumerate(headers):
        row = info_table.rows[r_idx]
        for c_idx, (text, is_bold) in enumerate(row_data):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
            run = p.add_run(text)
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9.5)
            if is_bold:
                run.font.bold = True
                run.font.color.rgb = theme_color
            else:
                run.font.bold = False
                
    # Add a thin line under the info table
    for r_idx in range(3):
        for cell in info_table.rows[r_idx].cells:
            # Set light border for info table
            set_cell_border(cell, 
                            top={"val": "none"}, 
                            bottom={"val": "none"}, 
                            left={"val": "none"}, 
                            right={"val": "none"})
            
    # Add a spacer
    spacer_p = doc.add_paragraph()
    spacer_p.paragraph_format.space_before = Pt(0)
    spacer_p.paragraph_format.space_after = Pt(4)
    
    # Create the main OMR Grid Table
    # 3 groups of questions:
    # Group 1: 1-50
    # Group 2: 51-100
    # Group 3: 101-150
    #
    # Each group has: Q# (0.35 in), 5 option columns (5 * 0.35 in = 1.75 in). Total = 2.1 in.
    # Group 1 (6 cols) + Spacer (1 col) + Group 2 (6 cols) + Spacer (1 col) + Group 3 (6 cols) = 20 columns.
    # Spacer column width: 0.35 in
    # Total width = 3 * 2.1 + 2 * 0.35 = 7.0 inches.
    
    num_rows = 50
    main_table = doc.add_table(rows=num_rows, cols=20)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_table.autofit = False
    
    # Column widths configuration
    col_widths = [
        # Group 1
        0.35, 0.35, 0.35, 0.35, 0.35, 0.35,
        # Spacer
        0.35,
        # Group 2
        0.35, 0.35, 0.35, 0.35, 0.35, 0.35,
        # Spacer
        0.35,
        # Group 3
        0.35, 0.35, 0.35, 0.35, 0.35, 0.35
    ]
    col_widths_inches = [Inches(w) for w in col_widths]
    
    group_starts = [0, 7, 14]
    group_ranges = [
        (1, 50),     # Group 1
        (51, 100),   # Group 2
        (101, 150)   # Group 3
    ]
    
    # Unicode Circled Letters for A, B, C, D, E
    circled_letters = ["Ⓐ", "Ⓑ", "Ⓒ", "Ⓓ", "Ⓔ"]
    
    border_style = {"sz": 12, "val": "single", "color": theme_hex}
    
    for r_idx in range(num_rows):
        row = main_table.rows[r_idx]
        # Set tight row height to ensure it fits on one page
        row.height = Pt(12.8)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        # Apply column widths to cells
        for c_idx, width in enumerate(col_widths_inches):
            cell = row.cells[c_idx]
            cell.width = width
            set_cell_margins(cell, top=10, bottom=10, left=5, right=5)
            # Remove all borders by default
            set_cell_border(cell, top={"val": "none"}, bottom={"val": "none"}, left={"val": "none"}, right={"val": "none"})

        # Populate the 3 groups
        for g_idx, start_col in enumerate(group_starts):
            q_start, q_end = group_ranges[g_idx]
            q_num = q_start + r_idx
            
            # Question number cell
            q_cell = row.cells[start_col]
            qp = q_cell.paragraphs[0]
            qp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            qp.paragraph_format.space_before = Pt(0)
            qp.paragraph_format.space_after = Pt(0)
            qp.paragraph_format.line_spacing = 1.0
            
            q_run = qp.add_run(f"{q_num}.")
            q_run.font.name = 'Segoe UI'
            q_run.font.size = Pt(8.5)
            q_run.font.bold = True
            q_run.font.color.rgb = theme_color
            
            # Option cells
            for opt_idx, letter in enumerate(circled_letters):
                opt_cell = row.cells[start_col + 1 + opt_idx]
                op = opt_cell.paragraphs[0]
                op.alignment = WD_ALIGN_PARAGRAPH.CENTER
                op.paragraph_format.space_before = Pt(0)
                op.paragraph_format.space_after = Pt(0)
                op.paragraph_format.line_spacing = 1.0
                
                o_run = op.add_run(letter)
                o_run.font.name = 'MS Gothic'
                o_run.font.size = Pt(11)
                o_run.font.bold = True
                o_run.font.color.rgb = theme_color
                
            # Apply vertical outer borders for each group to create clean vertical columns
            left_cell = row.cells[start_col]
            set_cell_border(left_cell, left=border_style)
            
            right_cell = row.cells[start_col + 5]
            set_cell_border(right_cell, right=border_style)
            
            # Top border for the first row of each group
            if r_idx == 0:
                for c_offset in range(6):
                    c_cell = row.cells[start_col + c_offset]
                    l_border = border_style if c_offset == 0 else {"val": "none"}
                    r_border = border_style if c_offset == 5 else {"val": "none"}
                    set_cell_border(c_cell, top=border_style, left=l_border, right=r_border)
                    
            # Bottom border for the last row of each group
            if r_idx == num_rows - 1:
                for c_offset in range(6):
                    c_cell = row.cells[start_col + c_offset]
                    l_border = border_style if c_offset == 0 else {"val": "none"}
                    r_border = border_style if c_offset == 5 else {"val": "none"}
                    set_cell_border(c_cell, bottom=border_style, left=l_border, right=r_border)

    # Word requires a paragraph after a table. Let's make sure it's tiny so it doesn't push to page 2.
    tail_p = doc.add_paragraph()
    tail_p.paragraph_format.space_before = Pt(0)
    tail_p.paragraph_format.space_after = Pt(0)
    tail_p.paragraph_format.line_spacing = 1.0
    tail_run = tail_p.add_run()
    tail_run.font.size = Pt(1)
    
    # Save document
    doc.save("d:\\Temp\\automation\\OMR_Sheet_5_Options.docx")
    print("OMR Sheet document created successfully!")

if __name__ == "__main__":
    create_omr_sheet()
