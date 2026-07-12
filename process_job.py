import os
import sys
import json
import sqlite3
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from fpdf.enums import XPos, YPos
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

# Dynamic data directory detection (Docker vs. Windows Host)
DATA_DIR = os.environ.get("DATA_DIR")
if not DATA_DIR:
    if os.path.exists("/data"):
        DATA_DIR = "/data"
    else:
        DATA_DIR = "D:/AI Job Automation"

# Standardize path separators
DATA_DIR = DATA_DIR.replace("\\", "/")

def get_path(subpath):
    return f"{DATA_DIR}/{subpath}"

# Ensure folders exist
def ensure_directories():
    directories = [
        get_path("Database"),
        get_path("Excel"),
        get_path("Resume"),
        get_path("TailoredResume"),
        get_path("CoverLetters"),
        get_path("Logs"),
        get_path("Config")
    ]
    for d in directories:
        os.makedirs(d, exist_ok=True)

# Logger function
def log_message(message, level="INFO"):
    ensure_directories()
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_file = get_path("Logs/automation.log")
    formatted = f"[{timestamp}] [{level}] {message}\n"
    with open(log_file, "a", encoding="utf-8") as f:
        f.write(formatted)
    print(formatted.strip())

# Initialize SQLite database
def init_db():
    db_path = get_path("Database/jobs.db")
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    # Candidate profile table
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS candidate_profile (
            id TEXT PRIMARY KEY,
            name TEXT,
            email TEXT,
            phone TEXT,
            skills TEXT,
            programming_languages TEXT,
            frameworks TEXT,
            libraries TEXT,
            databases TEXT,
            cloud_platforms TEXT,
            ai_skills TEXT,
            devops_skills TEXT,
            years_experience REAL,
            education TEXT,
            projects TEXT,
            certifications TEXT,
            keywords TEXT,
            desired_roles TEXT,
            github TEXT,
            linkedin TEXT,
            summary TEXT,
            updated_at TEXT
        )
    """)
    
    # Jobs table
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS jobs (
            job_url TEXT PRIMARY KEY,
            title TEXT,
            company TEXT,
            location TEXT,
            source TEXT,
            description TEXT,
            tags TEXT,
            match_score REAL,
            matched_skills TEXT,
            skill_gap TEXT,
            experience_match TEXT,
            location_match TEXT,
            remote_match TEXT,
            salary_match TEXT,
            priority_score REAL,
            recommendation TEXT,
            status TEXT,
            match_summary TEXT,
            tailored_bullets TEXT,
            cover_letter_opener TEXT,
            date_logged TEXT,
            contact_email TEXT,
            applied_date TEXT,
            follow_up_count INTEGER DEFAULT 0,
            next_follow_up_date TEXT
        )
    """)
    
    # Run migrations for existing databases to add missing columns dynamically
    for col, col_type in [("name", "TEXT"), ("email", "TEXT"), ("phone", "TEXT")]:
        try:
            cursor.execute(f"ALTER TABLE candidate_profile ADD COLUMN {col} {col_type}")
        except sqlite3.OperationalError:
            pass # Already exists
            
    for col, col_type in [
        ("contact_email", "TEXT"),
        ("applied_date", "TEXT"),
        ("follow_up_count", "INTEGER DEFAULT 0"),
        ("next_follow_up_date", "TEXT")
    ]:
        try:
            cursor.execute(f"ALTER TABLE jobs ADD COLUMN {col} {col_type}")
        except sqlite3.OperationalError:
            pass # Already exists
            
    conn.commit()
    conn.close()

# Save/Update Candidate Profile in DB
def save_candidate_profile(candidate):
    db_path = get_path("Database/jobs.db")
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    updated_at = datetime.datetime.now().isoformat()
    cursor.execute("""
        INSERT INTO candidate_profile (
            id, name, email, phone, skills, programming_languages, frameworks, libraries, databases,
            cloud_platforms, ai_skills, devops_skills, years_experience, education,
            projects, certifications, keywords, desired_roles, github, linkedin, summary, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(id) DO UPDATE SET
            name=excluded.name,
            email=excluded.email,
            phone=excluded.phone,
            skills=excluded.skills,
            programming_languages=excluded.programming_languages,
            frameworks=excluded.frameworks,
            libraries=excluded.libraries,
            databases=excluded.databases,
            cloud_platforms=excluded.cloud_platforms,
            ai_skills=excluded.ai_skills,
            devops_skills=excluded.devops_skills,
            years_experience=excluded.years_experience,
            education=excluded.education,
            projects=excluded.projects,
            certifications=excluded.certifications,
            keywords=excluded.keywords,
            desired_roles=excluded.desired_roles,
            github=excluded.github,
            linkedin=excluded.linkedin,
            summary=excluded.summary,
            updated_at=excluded.updated_at
    """, (
        "default_candidate",
        candidate.get("name", "Md Sadique Amin"),
        candidate.get("email", "mdsadiqueamin721786@gmail.com"),
        candidate.get("phone", "+91 9318302850"),
        json.dumps(candidate.get("skills", [])),
        json.dumps(candidate.get("programmingLanguages", [])),
        json.dumps(candidate.get("frameworks", [])),
        json.dumps(candidate.get("libraries", [])),
        json.dumps(candidate.get("databases", [])),
        json.dumps(candidate.get("cloudPlatforms", [])),
        json.dumps(candidate.get("aiSkills", [])),
        json.dumps(candidate.get("devOpsSkills", [])),
        float(candidate.get("yearsExperience", 0)),
        json.dumps(candidate.get("education", [])),
        json.dumps(candidate.get("topProjects", [])),
        json.dumps(candidate.get("certifications", [])),
        json.dumps(candidate.get("keywords", [])),
        json.dumps(candidate.get("targetTitles", [])),
        candidate.get("github", ""),
        candidate.get("linkedin", ""),
        candidate.get("summary", ""),
        updated_at
    ))
    
    conn.commit()
    conn.close()

# Save/Update Job in SQLite
def save_job_to_db(job, tailored):
    db_path = get_path("Database/jobs.db")
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    date_logged = datetime.datetime.now().strftime("%Y-%m-%d")
    
    cursor.execute("""
        INSERT INTO jobs (
            job_url, title, company, location, source, description, tags, match_score, matched_skills,
            skill_gap, experience_match, location_match, remote_match, salary_match, priority_score,
            recommendation, status, match_summary, tailored_bullets, cover_letter_opener, date_logged,
            contact_email, applied_date, follow_up_count, next_follow_up_date
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(job_url) DO UPDATE SET
            title=excluded.title,
            company=excluded.company,
            location=excluded.location,
            source=excluded.source,
            description=excluded.description,
            tags=excluded.tags,
            match_score=excluded.match_score,
            matched_skills=excluded.matched_skills,
            skill_gap=excluded.skill_gap,
            experience_match=excluded.experience_match,
            location_match=excluded.location_match,
            remote_match=excluded.remote_match,
            salary_match=excluded.salary_match,
            priority_score=excluded.priority_score,
            recommendation=excluded.recommendation,
            status=coalesce(jobs.status, excluded.status),
            match_summary=excluded.match_summary,
            tailored_bullets=excluded.tailored_bullets,
            cover_letter_opener=excluded.cover_letter_opener,
            contact_email=coalesce(jobs.contact_email, excluded.contact_email)
    """, (
        job.get("jobUrl", ""),
        job.get("title", ""),
        job.get("company", ""),
        job.get("jobLocation", ""),
        job.get("source", ""),
        job.get("description", ""),
        job.get("tags", ""),
        float(job.get("matchScore", 0)),
        job.get("matchedSkills", ""),
        tailored.get("skillGap", ""),
        tailored.get("experienceMatch", ""),
        tailored.get("locationMatch", ""),
        tailored.get("remoteMatch", ""),
        tailored.get("salaryMatch", "N/A"),
        float(tailored.get("priorityScore", job.get("matchScore", 0))),
        tailored.get("recommendation", ""),
        "Discovered",
        tailored.get("matchSummary", ""),
        " | ".join(tailored.get("tailoredBullets", [])) if isinstance(tailored.get("tailoredBullets"), list) else str(tailored.get("tailoredBullets", "")),
        tailored.get("coverLetterOpener", ""),
        date_logged,
        tailored.get("contactEmail", job.get("contactEmail", "")),
        None,
        0,
        None
    ))
    
    conn.commit()
    conn.close()

# Save/Update Job in Excel (JobTracker.xlsx)
def save_job_to_excel(job, tailored):
    excel_path = get_path("Excel/JobTracker.xlsx")
    
    # Check if exists, otherwise create
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        ws = wb.active
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Jobs CRM"
        # Style sheet
        ws.views.sheetView[0].showGridLines = True
        
        # Headers
        headers = [
            "Date", "Job Title", "Company", "Location", "Source", "Job URL", 
            "Match Score", "Matched Skills", "Priority Score", "Match Summary", 
            "Tailored Bullets", "Cover Letter Opener", "Status"
        ]
        ws.append(headers)
        
        # Style Header
        header_fill = PatternFill(start_color="1F497D", end_color="1F497D", fill_type="solid")
        header_font = Font(name="Segoe UI", size=11, bold=True, color="FFFFFF")
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            
    # Search if row exists based on Job URL
    row_to_write = None
    for r in range(2, ws.max_row + 1):
        val = ws.cell(row=r, column=6).value
        if val == job.get("jobUrl", ""):
            row_to_write = r
            break
            
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    bullets_str = " | ".join(tailored.get("tailoredBullets", [])) if isinstance(tailored.get("tailoredBullets"), list) else str(tailored.get("tailoredBullets", ""))
    
    row_data = [
        date_str,
        job.get("title", ""),
        job.get("company", ""),
        job.get("jobLocation", ""),
        job.get("source", ""),
        job.get("jobUrl", ""),
        float(job.get("matchScore", 0)),
        job.get("matchedSkills", ""),
        float(tailored.get("priorityScore", job.get("matchScore", 0))),
        tailored.get("matchSummary", ""),
        bullets_str,
        tailored.get("coverLetterOpener", ""),
        "Discovered"
    ]
    
    if row_to_write:
        for c, val in enumerate(row_data, 1):
            ws.cell(row=row_to_write, column=c, value=val)
    else:
        ws.append(row_data)
        row_to_write = ws.max_row
        
    # Apply styling
    thin_border = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )
    
    row_fill = PatternFill(start_color="F2F5F8" if row_to_write % 2 == 0 else "FFFFFF", end_color="F2F5F8" if row_to_write % 2 == 0 else "FFFFFF", fill_type="solid")
    row_font = Font(name="Segoe UI", size=10)
    
    for col_idx in range(1, len(row_data) + 1):
        cell = ws.cell(row=row_to_write, column=col_idx)
        cell.font = row_font
        cell.fill = row_fill
        cell.border = thin_border
        
        if col_idx in [1, 5, 7, 9, 13]:
            cell.alignment = Alignment(horizontal="center", vertical="top")
        elif col_idx == 6:
            cell.alignment = Alignment(horizontal="left", vertical="top")
            cell.hyperlink = cell.value
            cell.font = Font(name="Segoe UI", size=10, color="0000FF", underline="single")
        else:
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            val_str = str(cell.value or '')
            if cell.row == 1:
                max_len = max(max_len, len(val_str) + 4)
            else:
                max_len = max(max_len, min(len(val_str), 30))
        ws.column_dimensions[col_letter].width = max(max_len, 10)
        
    wb.save(excel_path)

# Helper to sanitize filename
def sanitize_filename(name):
    keep = [" ", "-", "_"]
    return "".join(c for c in name.strip() if c.isalnum() or c in keep).replace(" ", "_")

# Helper to sanitize string for standard Latin-1 PDF generation
def sanitize_for_pdf(text):
    if not text:
        return ""
    if not isinstance(text, str):
        text = str(text)
    
    # Replace common Unicode characters with Latin-1/ASCII equivalents
    replacements = {
        '\u2013': '-',  # en-dash
        '\u2014': '-',  # em-dash
        '\u2018': "'",  # left single quote
        '\u2019': "'",  # right single quote
        '\u201c': '"',  # left double quote
        '\u201d': '"',  # right double quote
        '\u2022': '*',  # bullet point
        '\u00a0': ' ',  # non-breaking space
        '\u200b': '',   # zero-width space
    }
    for key, val in replacements.items():
        text = text.replace(key, val)
        
    return text.encode('latin-1', 'replace').decode('latin-1')

# Generate DOCX Resume
def generate_docx_resume(candidate, job, tailored, docx_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    
    name = candidate.get("name") or "Md Sadique Amin"
    email = candidate.get("email") or "mdsadiqueamin721786@gmail.com"
    phone = candidate.get("phone") or "+91 9318302850"
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"{name.upper()}\n")
    run.font.size = Pt(18)
    run.bold = True
    
    contact_parts = []
    if email: contact_parts.append(f"Email: {email}")
    if phone: contact_parts.append(f"Phone: {phone}")
    if candidate.get("github"): contact_parts.append(f"Github: {candidate.get('github')}")
    if candidate.get("linkedin"): contact_parts.append(f"LinkedIn: {candidate.get('linkedin')}")
    contact_info = " | ".join(contact_parts) + "\n"
    title_p.add_run(contact_info).font.size = Pt(9.5)
    
    h = doc.add_paragraph()
    r = h.add_run("PROFESSIONAL SUMMARY")
    r.bold = True
    r.font.size = Pt(12)
    doc.add_paragraph(candidate.get("summary", ""))
    
    h = doc.add_paragraph()
    r = h.add_run(f"TAILORED HIGHLIGHTS (Targeting: {job.get('title')} at {job.get('company')})")
    r.bold = True
    r.font.size = Pt(12)
    
    bullets = tailored.get("tailoredBullets", [])
    if isinstance(bullets, str):
        bullets = [b.strip() for b in bullets.split("|") if b.strip()]
        
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
        
    h = doc.add_paragraph()
    r = h.add_run("CORE SKILLS")
    r.bold = True
    r.font.size = Pt(12)
    
    skills_categories = {
        "Languages": candidate.get("programmingLanguages", []),
        "Frameworks & Libraries": candidate.get("frameworks", []) + candidate.get("libraries", []),
        "Databases & Cloud": candidate.get("databases", []) + candidate.get("cloudPlatforms", []),
        "DevOps & AI Tools": candidate.get("devOpsSkills", []) + candidate.get("aiSkills", []),
    }
    
    for category, items in skills_categories.items():
        if items:
            p = doc.add_paragraph()
            r_cat = p.add_run(f"{category}: ")
            r_cat.bold = True
            p.add_run(", ".join(items))
            
    if candidate.get("topProjects"):
        h = doc.add_paragraph()
        r = h.add_run("RELEVANT PROJECTS")
        r.bold = True
        r.font.size = Pt(12)
        for proj in candidate.get("topProjects", []):
            doc.add_paragraph(proj, style='List Bullet')
            
    if candidate.get("certifications"):
        h = doc.add_paragraph()
        r = h.add_run("CERTIFICATIONS")
        r.bold = True
        r.font.size = Pt(12)
        p = doc.add_paragraph()
        p.add_run(", ".join(candidate.get("certifications", [])))
        
    if candidate.get("education"):
        h = doc.add_paragraph()
        r = h.add_run("EDUCATION")
        r.bold = True
        r.font.size = Pt(12)
        for edu in candidate.get("education", []):
            doc.add_paragraph(edu)
            
    doc.save(docx_path)

# Generate PDF Resume
class PDFResume(FPDF):
    def header(self):
        pass
    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", 0, align="C", new_x=XPos.RIGHT, new_y=YPos.TOP)

def generate_pdf_resume(candidate, job, tailored, pdf_path):
    pdf = PDFResume()
    pdf.add_page()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    
    name = sanitize_for_pdf(candidate.get("name") or "Md Sadique Amin")
    email = sanitize_for_pdf(candidate.get("email") or "mdsadiqueamin721786@gmail.com")
    phone = sanitize_for_pdf(candidate.get("phone") or "+91 9318302850")
    
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 10, name.upper(), 0, align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    
    pdf.set_font("Helvetica", "", 9)
    contact_parts = []
    if email: contact_parts.append(f"Email: {email}")
    if phone: contact_parts.append(f"Phone: {phone}")
    if candidate.get("github"): contact_parts.append(f"Github: {sanitize_for_pdf(candidate.get('github'))}")
    if candidate.get("linkedin"): contact_parts.append(f"LinkedIn: {sanitize_for_pdf(candidate.get('linkedin'))}")
    contact = " | ".join(contact_parts)
    pdf.cell(0, 5, contact, 0, align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(5)
    
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(31, 73, 125)
    pdf.cell(0, 6, "PROFESSIONAL SUMMARY", 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0, 0, 0)
    pdf.multi_cell(0, 5, sanitize_for_pdf(candidate.get("summary", "")))
    pdf.ln(4)
    
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(31, 73, 125)
    pdf.cell(0, 6, sanitize_for_pdf(f"TAILORED MATCH HIGHLIGHTS ({job.get('company')} - {job.get('title')})"), 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0, 0, 0)
    bullets = tailored.get("tailoredBullets", [])
    if isinstance(bullets, str):
        bullets = [b.strip() for b in bullets.split("|") if b.strip()]
    for bullet in bullets:
        pdf.cell(5, 5, chr(149), 0, new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.multi_cell(0, 5, sanitize_for_pdf(bullet))
        pdf.ln(1)
    pdf.ln(3)
    
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(31, 73, 125)
    pdf.cell(0, 6, "TECHNICAL SKILLS", 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0, 0, 0)
    
    skills_map = [
        ("Languages", candidate.get("programmingLanguages", [])),
        ("Frameworks", candidate.get("frameworks", []) + candidate.get("libraries", [])),
        ("Databases/Cloud", candidate.get("databases", []) + candidate.get("cloudPlatforms", [])),
        ("DevOps/AI", candidate.get("devOpsSkills", []) + candidate.get("aiSkills", []))
    ]
    for cat_name, items in skills_map:
        if items:
            pdf.set_font("Helvetica", "B", 10)
            pdf.write(5, sanitize_for_pdf(f"{cat_name}: "))
            pdf.set_font("Helvetica", "", 10)
            pdf.write(5, sanitize_for_pdf(", ".join(items)) + "\n")
    pdf.ln(4)
    
    if candidate.get("topProjects"):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(31, 73, 125)
        pdf.cell(0, 6, "SELECTED PROJECTS", 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(0, 0, 0)
        for proj in candidate.get("topProjects", []):
            pdf.cell(5, 5, chr(149), 0, new_x=XPos.RIGHT, new_y=YPos.TOP)
            pdf.multi_cell(0, 5, sanitize_for_pdf(proj))
            pdf.ln(1)
        pdf.ln(3)
        
    if candidate.get("education") or candidate.get("certifications"):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(31, 73, 125)
        pdf.cell(0, 6, "EDUCATION & CREDENTIALS", 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(0, 0, 0)
        
        if candidate.get("education"):
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 5, "Education:", 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 10)
            for edu in candidate.get("education", []):
                pdf.cell(0, 5, sanitize_for_pdf(edu), 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)
            
        if candidate.get("certifications"):
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 5, "Certifications:", 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, sanitize_for_pdf(", ".join(candidate.get("certifications", []))))
            
    pdf.output(pdf_path)

# Generate DOCX Cover Letter
def generate_docx_cover_letter(candidate, job, tailored, docx_path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    name = candidate.get("name") or "Md Sadique Amin"
    email = candidate.get("email") or "mdsadiqueamin721786@gmail.com"
    phone = candidate.get("phone") or "+91 9318302850"
    
    p = doc.add_paragraph()
    r = p.add_run(f"{name.upper()}\n")
    r.bold = True
    r.font.size = Pt(16)
    
    contact_info = f"{email} | {phone}\n\n" if phone else f"{email}\n\n"
    p.add_run(contact_info)
    
    date_str = datetime.datetime.now().strftime("%B %d, %Y")
    p.add_run(f"{date_str}\n\n")
    p.add_run(f"Hiring Team\n{job.get('company')}\n\n")
    
    doc.add_paragraph(f"Dear Hiring Team at {job.get('company')},")
    doc.add_paragraph(tailored.get("coverLetterOpener", ""))
    
    body_text = (
        f"I am writing to express my strong interest in the {job.get('title')} opportunity. "
        f"As an automation engineer with experience in fields like {', '.join(candidate.get('skills', [])[:5])}, "
        f"I have successfully delivered projects utilizing tools like {', '.join(candidate.get('programmingLanguages', [])[:3])}. "
        f"I believe my background aligns closely with the objectives of your team."
    )
    doc.add_paragraph(body_text)
    doc.add_paragraph("Thank you for your time and consideration. I look forward to discussing how my experience can contribute to your goals.")
    doc.add_paragraph(f"\nSincerely,\n\n{name.upper()}")
    doc.save(docx_path)
 
# Generate PDF Cover Letter
def generate_pdf_cover_letter(candidate, job, tailored, pdf_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=20)
    
    name = sanitize_for_pdf(candidate.get("name") or "Md Sadique Amin")
    email = sanitize_for_pdf(candidate.get("email") or "mdsadiqueamin721786@gmail.com")
    phone = sanitize_for_pdf(candidate.get("phone") or "+91 9318302850")
    
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 8, name.upper(), 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 10)
    
    contact_info = f"Email: {email} | Phone: {phone}" if phone else f"Email: {email}"
    pdf.cell(0, 5, contact_info, 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(8)
    
    date_str = datetime.datetime.now().strftime("%B %d, %Y")
    pdf.cell(0, 5, date_str, 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)
    pdf.cell(0, 5, sanitize_for_pdf(f"Hiring Team at {job.get('company')}"), 0, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(6)
    
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, sanitize_for_pdf(f"Dear Hiring Team at {job.get('company')},"))
    pdf.ln(4)
    pdf.multi_cell(0, 6, sanitize_for_pdf(tailored.get("coverLetterOpener", "")))
    pdf.ln(4)
    
    body_text = (
        f"I am writing to express my strong interest in the {job.get('title')} opportunity. "
        f"As an automation engineer with experience in fields like {', '.join(candidate.get('skills', [])[:5])}, "
        f"I have successfully delivered projects utilizing tools like {', '.join(candidate.get('programmingLanguages', [])[:3])}. "
        f"I believe my background aligns closely with the objectives of your team."
    )
    pdf.multi_cell(0, 6, sanitize_for_pdf(body_text))
    pdf.ln(4)
    pdf.multi_cell(0, 6, "Thank you for your time and consideration. I look forward to discussing how my experience can contribute to your goals.")
    pdf.ln(8)
    pdf.multi_cell(0, 6, f"Sincerely,\n\n{name.upper()}")
    
    pdf.output(pdf_path)

# Main processor flow
def main():
    ensure_directories()
    init_db()
    
    try:
        payload = json.load(sys.stdin)
        candidate = payload.get("candidate", {})
        job = payload.get("job", {})
        
        if candidate:
            save_candidate_profile(candidate)
            
        if not job:
            log_message("Candidate profile updated, no job payload provided.", "INFO")
            return
            
        company_clean = sanitize_filename(job.get("company", "Unknown"))
        title_clean = sanitize_filename(job.get("title", "Unknown"))
        
        tailored = job.get("tailored", {})
        if not tailored:
            tailored = {
                "matchSummary": job.get("matchSummary", ""),
                "tailoredBullets": job.get("tailoredBullets", []),
                "coverLetterOpener": job.get("coverLetterOpener", ""),
                "skillGap": job.get("skillGap", ""),
                "experienceMatch": job.get("experienceMatch", ""),
                "locationMatch": job.get("locationMatch", ""),
                "remoteMatch": job.get("remoteMatch", ""),
                "salaryMatch": job.get("salaryMatch", "N/A"),
                "priorityScore": job.get("priorityScore", job.get("matchScore", 0)),
                "recommendation": job.get("recommendation", "")
            }
            
        # 1. SQLite Logging
        save_job_to_db(job, tailored)
        
        # 2. Excel Logging
        save_job_to_excel(job, tailored)
        
        # 3. File Names
        resume_docx = get_path(f"TailoredResume/Resume_{company_clean}_{title_clean}.docx")
        resume_pdf = get_path(f"TailoredResume/Resume_{company_clean}_{title_clean}.pdf")
        cl_docx = get_path(f"CoverLetters/CoverLetter_{company_clean}_{title_clean}.docx")
        cl_pdf = get_path(f"CoverLetters/CoverLetter_{company_clean}_{title_clean}.pdf")
        email_txt = get_path(f"CoverLetters/ColdEmail_{company_clean}_{title_clean}.txt")
        
        # 4. Tailored Resume DOCX & PDF
        generate_docx_resume(candidate, job, tailored, resume_docx)
        generate_pdf_resume(candidate, job, tailored, resume_pdf)
        
        # 5. Cover Letter DOCX & PDF
        generate_docx_cover_letter(candidate, job, tailored, cl_docx)
        generate_pdf_cover_letter(candidate, job, tailored, cl_pdf)
        
        # 6. Cold Email TXT
        name = candidate.get("name") or "Md Sadique Amin"
        email = candidate.get("email") or "mdsadiqueamin721786@gmail.com"
        email_content = (
            f"Subject: Application for {job.get('title')} - {name.upper()}\n\n"
            f"Dear Hiring Team at {job.get('company')},\n\n"
            f"{tailored.get('coverLetterOpener', '')}\n\n"
            f"My profile matches your requirements for {job.get('title')}. "
            f"Please find attached my tailored resume and cover letter.\n\n"
            f"Best regards,\n{name.upper()}\n{email}"
        )
        with open(email_txt, "w", encoding="utf-8") as f:
            f.write(email_content)
            
        log_message(f"Successfully processed and generated files for {job.get('title')} at {job.get('company')}.", "INFO")
        
        print(json.dumps({
            "status": "success",
            "company": job.get("company"),
            "title": job.get("title"),
            "db_path": get_path("Database/jobs.db"),
            "excel_path": get_path("Excel/JobTracker.xlsx"),
            "resume_pdf": resume_pdf,
            "cl_pdf": cl_pdf,
            "cold_email_txt": email_txt
        }))
        
    except Exception as e:
        log_message(f"Failed to process job: {str(e)}", "ERROR")
        print(json.dumps({
            "status": "error",
            "message": str(e)
        }))
        sys.exit(1)

if __name__ == "__main__":
    main()
